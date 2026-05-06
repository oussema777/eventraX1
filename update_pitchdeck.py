#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Apply the 3 comment changes to the pitch deck document."""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

doc = Document('Eventra_PitchDeck_Content_SoftLanding2026.docx')

# ============================================
# COMMENT #1: Enhance Slide 5 - Custom Forms & Evaluation
# ============================================

for i, para in enumerate(doc.paragraphs):
    # Add Custom Forms bullet in PRE-EVENT (after B2B Matchmaking Setup)
    if 'B2B Matchmaking Setup' in para.text and 'interest matching' in para.text:
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            if 'DURING EVENT' in doc.paragraphs[j].text:
                if not doc.paragraphs[j-1].text.strip():
                    target = doc.paragraphs[j-1]
                    target.style = doc.paragraphs[i].style
                    target.text = (
                        'Custom Forms & Evaluation \u2014 Pre-event surveys, registration questions, '
                        'speaker/session evaluation forms. Replaces Google Forms entirely for '
                        'masterclasses, workshops, and conferences'
                    )
                    print(f'Added Custom Forms bullet in PRE-EVENT at paragraph {j-1}')
                break
        break

# Update Analytics bullet to include session ratings
for i, para in enumerate(doc.paragraphs):
    if 'Analytics & ROI Reporting' in para.text and 'Attendance funnels' in para.text:
        para.text = 'Analytics & ROI Reporting \u2014 Attendance funnels, revenue, sponsor ROI, session ratings'
        print(f'Enhanced Analytics bullet at {i}')
        break

# Update Follow-up Campaigns to highlight evaluation forms
for i, para in enumerate(doc.paragraphs):
    if 'Follow-up Campaigns' in para.text and 'surveys' in para.text:
        para.text = (
            'Post-Event Evaluation & Follow-up \u2014 Real-time session evaluation forms, '
            'attendee satisfaction surveys, thank-you emails, NPS collection, next-event '
            'promotion. Built-in forms replace Google Forms for workshops, masterclasses, '
            'and conferences'
        )
        print(f'Enhanced Follow-up bullet at {i}')
        break

# Update Slide 5 visual direction
for i, para in enumerate(doc.paragraphs):
    if 'Three-column layout' in para.text and 'Pre' in para.text and 'During' in para.text:
        para.text = (
            'Three-column layout (Pre \u2192 During \u2192 Post) with connecting arrow flow. '
            '4-5 features per column with icons. Custom Forms highlighted as a cross-phase '
            'feature spanning Pre and Post. Actual Eventra UI screenshots embedded subtly '
            'behind each phase. Clean, not crowded.'
        )
        print(f'Updated Slide 5 visual at {i}')
        break

# Update Slide 5 speaker notes
for i, para in enumerate(doc.paragraphs):
    if 'Eventra is not a website builder or a ticketing tool' in para.text:
        para.text = (
            'Eventra is not a website builder or a ticketing tool \u2014 it is a complete event '
            'engine. Before the event: build your entire online presence with 15-plus drag-and-drop '
            'blocks, set up ticketing, launch email campaigns, configure matchmaking, and create '
            'custom registration and evaluation forms that replace Google Forms entirely. During: '
            'QR check-in, live AI-powered B2B meetings, trade mission coordination, real-time '
            'dashboard. After: full analytics, real-time session evaluation for masterclasses and '
            'workshops, attendee satisfaction surveys, data export, follow-up emails, and one-click '
            'event duplication. The custom forms system is used in both pre-event and post-event '
            'phases \u2014 from registration questions to real-time workshop evaluations. Everything '
            'that used to require 5 separate tools and 5 separate subscriptions now lives in one '
            'platform, accessible in 3 languages.'
        )
        print(f'Updated Slide 5 speaker notes at {i}')
        break

# ============================================
# COMMENT #0: Expand Slide 4 - 15 Competitor Comparison
# ============================================

# Update comparison header
for i, para in enumerate(doc.paragraphs):
    if 'TOP 5 COMPETITOR COMPARISON' in para.text:
        para.text = 'COMPETITOR COMPARISON \u2014 15 Platforms Benchmarked'
        print(f'Updated comparison header at {i}')
        break

# Update positioning map quadrants
for i, para in enumerate(doc.paragraphs):
    if 'Bottom-Left' in para.text and 'Eventbrite' in para.text:
        para.text = 'Bottom-Left (Low price, Low features): Eventbrite, Ticket Tailor, Splash, Accelevents'
        print(f'Updated positioning Bottom-Left at {i}')
    elif 'Top-Right' in para.text and 'Cvent' in para.text:
        para.text = 'Top-Right (High price, High features): Cvent, Bizzabo, Hubilo, Whova, vFairs'
        print(f'Updated positioning Top-Right at {i}')
    elif 'Top-Left' in para.text and 'EVENTRA' in para.text:
        para.text = 'Top-Left (Low price, High features): EVENTRA \u2190 This is the white space'
        print(f'Updated positioning Top-Left at {i}')
    elif 'Bottom-Right' in para.text and 'Hopin' in para.text:
        para.text = 'Bottom-Right (High price, Low features): Hopin/RingCentral, Airmeet, Zuddl (virtual-only or narrow focus)'
        print(f'Updated positioning Bottom-Right at {i}')

# Add full 15-competitor feature comparison table
for i, para in enumerate(doc.paragraphs):
    if 'Our Unfair Advantages' in para.text:
        if not doc.paragraphs[i-1].text.strip():
            target = doc.paragraphs[i-1]
            target.text = (
                'Feature Comparison (Companies across top, Features as rows):\n\n'
                'COMPETITORS: Eventra | Eventbrite | Cvent | Bizzabo | Hubilo | Whova | Splash | Hopin | Airmeet | vFairs | Swoogo | Accelevents | Zuddl | Grip | Brella\n\n'
                'FEATURES (Eventra has ALL):\n'
                '1. No-Code Website Builder: Eventra \u2713 | Eventbrite \u2717 | Cvent \u2713 | Bizzabo \u2713 | Hubilo \u2713 | Whova \u2717 | Splash \u2713 | Hopin \u2717 | Airmeet \u2717 | vFairs \u2713 | Swoogo \u2713 | Accelevents \u2717 | Zuddl \u2717 | Grip \u2717 | Brella \u2717\n'
                '2. Ticketing & Registration: Eventra \u2713 | Eventbrite \u2713 | Cvent \u2713 | Bizzabo \u2713 | Hubilo \u2713 | Whova \u2713 | Splash \u2713 | Hopin \u2713 | Airmeet \u2713 | vFairs \u2713 | Swoogo \u2713 | Accelevents \u2713 | Zuddl \u2713 | Grip \u2717 | Brella \u2717\n'
                '3. Email Marketing Campaigns: Eventra \u2713 | Eventbrite Basic | Cvent \u2713 | Bizzabo \u2713 | Hubilo \u2713 | Whova \u2717 | Splash \u2713 | Hopin \u2717 | Airmeet \u2717 | vFairs \u2717 | Swoogo \u2713 | Accelevents \u2717 | Zuddl \u2717 | Grip \u2717 | Brella \u2717\n'
                '4. B2B Matchmaking: Eventra \u2713 | Eventbrite \u2717 | Cvent \u2713 | Bizzabo \u2713 | Hubilo \u2713 | Whova \u2713 | Splash \u2717 | Hopin \u2717 | Airmeet \u2717 | vFairs \u2717 | Swoogo \u2717 | Accelevents \u2717 | Zuddl \u2717 | Grip \u2713 | Brella \u2713\n'
                '5. On-Site QR Check-In: Eventra \u2713 | Eventbrite \u2713 | Cvent \u2713 | Bizzabo \u2713 | Hubilo \u2713 | Whova \u2713 | Splash \u2717 | Hopin \u2717 | Airmeet \u2717 | vFairs \u2717 | Swoogo \u2713 | Accelevents \u2713 | Zuddl \u2717 | Grip \u2717 | Brella \u2717\n'
                '6. Custom Evaluation Forms: Eventra \u2713 | Eventbrite \u2717 | Cvent \u2713 | Bizzabo \u2713 | Hubilo \u2717 | Whova \u2713 | Splash \u2717 | Hopin \u2717 | Airmeet \u2717 | vFairs \u2717 | Swoogo \u2717 | Accelevents \u2717 | Zuddl \u2717 | Grip \u2717 | Brella \u2717\n'
                '7. Speaker & Agenda Management: Eventra \u2713 | Eventbrite \u2717 | Cvent \u2713 | Bizzabo \u2713 | Hubilo \u2713 | Whova \u2713 | Splash \u2717 | Hopin \u2713 | Airmeet \u2713 | vFairs \u2713 | Swoogo \u2713 | Accelevents \u2713 | Zuddl \u2713 | Grip \u2717 | Brella \u2717\n'
                '8. Sponsor Management: Eventra \u2713 | Eventbrite \u2717 | Cvent \u2713 | Bizzabo \u2713 | Hubilo \u2713 | Whova \u2713 | Splash \u2717 | Hopin \u2713 | Airmeet \u2717 | vFairs \u2713 | Swoogo \u2717 | Accelevents \u2717 | Zuddl \u2717 | Grip \u2717 | Brella \u2717\n'
                '9. Trade Mission Tools: Eventra \u2713 | ALL OTHERS \u2717 (Unique to Eventra)\n'
                '10. Trilingual (EN/FR/AR + RTL): Eventra \u2713 | ALL OTHERS \u2717 (Unique to Eventra)\n'
                '11. Analytics & ROI Dashboard: Eventra \u2713 | Eventbrite Basic | Cvent \u2713 | Bizzabo \u2713 | Hubilo \u2713 | Whova \u2713 | Splash \u2713 | Hopin \u2713 | Airmeet \u2713 | vFairs \u2713 | Swoogo \u2713 | Accelevents \u2713 | Zuddl \u2713 | Grip \u2713 | Brella \u2713\n'
                '12. Data Export (CSV/Excel): Eventra \u2713 | Eventbrite \u2713 | Cvent \u2713 | Bizzabo \u2713 | Hubilo \u2713 | Whova \u2713 | Splash \u2713 | Hopin \u2713 | Airmeet \u2713 | vFairs \u2713 | Swoogo \u2713 | Accelevents \u2713 | Zuddl \u2713 | Grip \u2713 | Brella \u2713\n'
                '13. Event Duplication: Eventra \u2713 | Eventbrite \u2713 | Cvent \u2713 | Bizzabo \u2713 | Hubilo \u2717 | Whova \u2717 | Splash \u2713 | Hopin \u2717 | Airmeet \u2717 | vFairs \u2717 | Swoogo \u2713 | Accelevents \u2717 | Zuddl \u2717 | Grip \u2717 | Brella \u2717\n'
                '14. Access Code / Private Events: Eventra \u2713 | Eventbrite \u2713 | Cvent \u2713 | Bizzabo \u2713 | Hubilo \u2717 | Whova \u2717 | Splash \u2717 | Hopin \u2717 | Airmeet \u2717 | vFairs \u2717 | Swoogo \u2717 | Accelevents \u2717 | Zuddl \u2717 | Grip \u2717 | Brella \u2717\n'
                '15. Under \u20ac500/year: Eventra \u2713 (\u20ac432) | Eventbrite \u2713 | Cvent \u2717 (\u20ac18K) | Bizzabo \u2717 (\u20ac16.5K) | Hubilo \u2717 (\u20ac12K) | Whova \u2717 (\u20ac5K+) | Splash \u2717 | Hopin \u2717 | Airmeet \u2717 | vFairs \u2717 | Swoogo \u2717 | Accelevents \u2717 | Zuddl \u2717 | Grip \u2717 (\u20ac10K+) | Brella \u2717 (\u20ac10K+)\n\n'
                'EVENTRA SCORE: 15/15 features | Next closest: Cvent (12/15) at 38x the price'
            )
            print(f'Added full 15-competitor comparison table at {i-1}')
        break

# Update appendix reference
for i, para in enumerate(doc.paragraphs):
    if 'Full 10-competitor analysis' in para.text:
        para.text = 'Full 15-competitor analysis available in appendix.'
        print(f'Updated appendix reference at {i}')
        break

# Update appendix header
for i, para in enumerate(doc.paragraphs):
    if 'Full 10-Competitor Comparison Table' in para.text:
        para.text = 'A. Full 15-Competitor Comparison Table'
        print(f'Updated appendix header at {i}')
        break

# Update Slide 4 visual direction
for i, para in enumerate(doc.paragraphs):
    if 'positioning matrix' in para.text and 'competitor logos' in para.text:
        para.text = (
            '2\u00d72 positioning matrix with 15 competitor logos plotted. Eventra in the '
            '"white space" top-left quadrant highlighted in blue. Below: horizontal comparison '
            'table \u2014 competitors across the top, 15 features as rows. Eventra column '
            'fully checked (\u2713). Clean, scannable, data-dense.'
        )
        print(f'Updated Slide 4 visual at {i}')
        break

# Update Slide 4 speaker notes
for i, para in enumerate(doc.paragraphs):
    if 'Here is where Eventra sits in the competitive landscape' in para.text:
        para.text = (
            'Here is where Eventra sits in the competitive landscape. We benchmarked 15 platforms '
            'across 15 features. The market has two extremes: cheap but limited tools like Eventbrite '
            '\u2014 just ticketing, no builder, no matchmaking \u2014 and full-featured enterprise '
            'platforms like Cvent and Bizzabo at 16 to 18 thousand euros per year. The white space is '
            'top-left: full features at accessible pricing. That is exactly where Eventra sits. We '
            'score 15 out of 15 on our feature comparison. The next closest competitor is Cvent at '
            '12 out of 15 \u2014 but at 38 times our price. And we have two features no competitor '
            'offers at any price: native Arabic with full RTL support for 400 million speakers, and '
            'purpose-built trade mission tools. Our cost advantage is structural: lean team, no '
            'investor bloat, and emerging-market development efficiency.'
        )
        print(f'Updated Slide 4 speaker notes at {i}')
        break

# ============================================
# COMMENT #2: Add payment gateway to Slide 12 - The Ask
# ============================================

for i, para in enumerate(doc.paragraphs):
    if 'Mentorship on European market entry' in para.text:
        para.text = (
            'Mentorship on European market entry, GDPR compliance, payment gateway setup '
            '(Stripe/payment integration for EU transactions), and commercial strategy'
        )
        print(f'Updated mentorship bullet with payment gateway at {i}')
        break

# Update Slide 12 speaker notes
for i, para in enumerate(doc.paragraphs):
    if 'Help us set up our European entity' in para.text and 'Mentor us on European market entry' in para.text:
        old = 'Mentor us on European market entry and GDPR.'
        new = (
            'Mentor us on European market entry, GDPR, and payment gateway setup '
            '\u2014 we need help integrating Stripe or a European payment processor '
            'for ticket transactions.'
        )
        para.text = para.text.replace(old, new)
        print(f'Updated Slide 12 speaker notes at {i}')
        break

# ============================================
# SAVE
# ============================================
doc.save('Eventra_PitchDeck_Content_SoftLanding2026_v2.docx')
print('\n=== ALL 3 COMMENTS APPLIED AND SAVED ===')
