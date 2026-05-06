import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

style_h1 = doc.styles['Heading 1']
style_h1.font.size = Pt(20)
style_h1.font.color.rgb = RGBColor(0x06, 0x84, 0xF5)
style_h1.font.bold = True

# Helper functions
def add_question(num, text, required=True):
    p = doc.add_paragraph()
    run = p.add_run(f'Q{num}. {text}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x0B, 0x26, 0x41)
    if required:
        req = p.add_run(' *')
        req.font.color.rgb = RGBColor(0xE0, 0x3E, 0x3E)
        req.bold = True
    p.paragraph_format.space_before = Pt(16)
    return p

def add_label(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_answer_block(lines):
    """Add answer as properly formatted paragraphs"""
    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(4)

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

def add_fill_in(text):
    p = doc.add_paragraph()
    run = p.add_run(f'\u2B1C TO FILL: {text}')
    run.bold = True
    run.font.color.rgb = RGBColor(0xE0, 0x7A, 0x00)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(12)
    return p

# ====== DOCUMENT START ======

title = doc.add_heading('SOFT LANDING PROVENCE AFRICA CONNECT 2026', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Application Form \u2014 Eventra by SEKETAK Solutions')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Deadline: April 23, 2026  |  Submit via F6S  |  Demo: eventra.cloud')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
info.paragraph_format.space_after = Pt(24)

# ── Q1 ──
add_question(1, 'Which of our themes does your project match?')
add_label('(GreenTech, HealthTech, BlueTech, AgriTech, Cultural and Creative Industries)')
add_answer_block(['Cultural and Creative Industries'])

# ── Q2 ──
add_question(2, 'How did you hear about our Program?')
add_label('(I received an email, From a partner organization, Social networks, From the press)')
add_answer_block(['From a partner organization'])

# ── Q3 ──
add_question(3, 'What are your main objectives in joining our program?')
add_answer_block([
    'Our primary objective is to establish Eventra\'s European headquarters in Aix-Marseille-Provence and position the platform as the leading event management solution bridging Africa and Europe. Specifically, we aim to:',
    '',
    '1. Establish a legal European entity to serve French and EU clients with local invoicing, GDPR compliance, and euro-denominated pricing \u2014 unlocking the \u20AC30B+ European event tech market.',
    '',
    '2. Build strategic B2B partnerships with event organizers, trade bodies, chambers of commerce, and institutional actors across the Mediterranean basin \u2014 leveraging Marseille\'s unique position as a Euro-African gateway.',
    '',
    '3. Accelerate commercialization by transitioning from our validated product (4 international events hosted, 1,200+ registered users) to a recurring revenue SaaS model, supported by the program\'s investor network and fundraising guidance.',
    '',
    '4. Scale our Africa-Europe bridge \u2014 Eventra is already natively trilingual (English, French, Arabic) and battle-tested across Egypt, C\u00f4te d\'Ivoire, and Tunisia. Aix-Marseille-Provence is the natural hub to connect both continents\' event ecosystems.',
])

# ── Q4 ──
add_question(4, 'Explain your project in a few words (50 max.) to a 7-year-old child or a 77-year-old grandmother.')
add_label('(50 words maximum)')
add_answer_block([
    'Eventra is like a magic website builder, but for events. You click, design your event page, sell tickets, match people who should meet each other, and manage everything \u2014 all in one place. It speaks French, English, and Arabic so everyone in Africa and Europe can use it.'
])

# ── Q5 ──
add_question(5, 'Can you precise your innovation and how is your solution innovative compared to existing ones?')
add_label('(technology, business model, market, impact...)')
add_answer_block([
    'Eventra\'s innovation sits at the intersection of four dimensions:',
    '',
    '1. Design Studio \u2014 No-Code Event Page Builder',
    'Unlike Eventbrite or Meetup that offer rigid templates, Eventra provides a drag-and-drop Design Studio with 15+ customizable blocks (hero, speakers, agenda, sponsors, gallery, FAQ, countdown, etc.). Event organizers build fully branded landing pages without designers or developers \u2014 comparable to what Wix did for websites, but purpose-built for events.',
    '',
    '2. AI-Powered B2B Matchmaking',
    'Our intelligent matchmaking engine connects event attendees based on their profiles, interests, and business objectives. This goes beyond simple networking \u2014 it creates structured, high-value B2B meetings within events. No competitor in the African market offers this natively integrated.',
    '',
    '3. Trilingual by Architecture, Not Afterthought',
    'Eventra is built from the ground up in English, French, and Arabic (including RTL support). This isn\'t a translation layer \u2014 it\'s core architecture. This makes us uniquely positioned for the Africa-Europe-MENA corridor where these three languages dominate business communication.',
    '',
    '4. All-in-One African Event Infrastructure',
    'While Western platforms like Hopin, Eventbrite, and Bizzabo focus on mature markets, they don\'t address Africa\'s specific needs: mixed payment methods, multilingual audiences, B2B trade missions, and institutional events. Eventra is the first event tech platform built by Africans, for the Africa-Europe corridor, with enterprise features: custom registration forms, attendee management, enterprise email campaigns, ticketing, speaker management, and session scheduling \u2014 all unified.',
    '',
    'Business Model Innovation: We don\'t just digitize events \u2014 we monetize the connections within them. Our B2B matchmaking creates measurable ROI for trade missions and investment forums, a model validated through our partnership with GITS (Global Investment and Trade Summit).',
])

# ── Q6 ──
add_question(6, 'Explain your Business Model (few lines)')
add_answer_block([
    'Eventra operates a tiered model combining subscriptions and event-based revenue:',
    '',
    '\u2022 Free tier: Basic event creation for small events \u2014 drives adoption and lets organizers experience the platform.',
    '\u2022 Paid plans: Monthly or per-event subscriptions unlocking the full Design Studio, custom branding, analytics, B2B matchmaking, and enterprise email campaigns. Pricing is being validated with early adopters and will be finalized during our commercial launch phase.',
    '\u2022 Enterprise packages: Custom pricing for large-scale institutional events (trade missions, investment summits, conferences) \u2014 includes dedicated setup, white-label options, and on-site support.',
    '\u2022 Transaction fee: A percentage commission on paid ticket sales processed through the platform.',
    '',
    'Our go-to-market targets institutional event organizers (chambers of commerce, trade bodies, ministries) who run recurring annual programs \u2014 ensuring predictable, high-value contracts. We are currently in pilot phase, validating willingness-to-pay with our existing network before locking in final pricing.',
])

# ── Q7 ──
add_question(7, 'Describe the founding team (Full name / % stake / Role in the company)')
add_answer_block([
    'Assem Kamel / [__]% / CEO & Co-Founder \u2014 B2B strategy, partnerships, and business development. Serial entrepreneur, co-founder of GITS (Global Investment and Trade Summit) and The LEE Experience (Lebanon). Certified Business Coach (ILO). 15+ years in international trade, investment facilitation, and cross-border business development across MENA and Africa.',
    '',
    'Douja Gharbi / [__]% / Co-Founder & Strategic Partnerships \u2014 CEO of RedStart Tunisia, one of North Africa\'s leading startup accelerators. 30+ years in entrepreneurship and economic development. Supported 800+ projects and hundreds of entrepreneurs. ESIL Angel for Tunisia. Expert in incubation, acceleration, and Euro-African ecosystem building.',
    '',
    'Oussama Lamine / [__]% / CTO & Co-Founder \u2014 Full-stack developer and product architect. Built Eventra end-to-end: React/TypeScript frontend, Supabase backend, AI integrations, and the Design Studio engine. Responsible for all technical development, product roadmap, and platform infrastructure.',
])
add_fill_in('Replace [__]% with exact stake percentages for each founder')

# ── Q8 ──
add_question(8, 'How is the team currently founded (technical, sales, marketing, development...)?')
add_answer_block([
    'The team is structured across three complementary pillars:',
    '',
    '\u2022 Technology & Product (Oussama Lamine): Full-stack development, platform architecture, AI/ML integrations, DevOps, and product design. Currently sole technical lead \u2014 building the engineering team is a key program objective.',
    '',
    '\u2022 Business Development & Sales (Assem Kamel): B2B partnerships, client acquisition, trade mission relationships, investor relations, and international expansion strategy. Leverages his GITS network spanning Egypt, Lebanon, C\u00f4te d\'Ivoire, and the broader MENA-Africa corridor.',
    '',
    '\u2022 Ecosystem & Strategic Partnerships (Douja Gharbi): Institutional partnerships, accelerator/incubator network, Euro-African corridor development, program design, and fundraising strategy. Brings RedStart Tunisia\'s network of 800+ supported startups, European institutional partners, and access to INVESTMED/UNIDO programs.',
    '',
    'Current gap we aim to fill through the program: A dedicated European commercial presence (sales + marketing) based in Aix-Marseille-Provence to drive client acquisition in the French and EU markets.',
])

# ── Q9 ──
add_question(9, 'What stage has your project reached?')
add_label('(Creation/Launch, Access to your market, Proven market traction, Development)')
add_answer_block([
    'Access to your market (MVP, early adopters/customers...)',
    '',
    'Eventra is a fully functional MVP that has been battle-tested across 4 international events in 3 African countries. We have 1,200+ registered user accounts, proven product-market validation, and are now transitioning from pilot phase to commercial launch.',
])

# ── Q10 ──
add_question(10, 'Sales in year N and N-1')
add_answer_block([
    'Year N (2026): \u20AC0 \u2014 Pre-revenue. Platform has been deployed free-of-charge across pilot events to validate product-market fit and build our user base.',
    '',
    'Year N-1 (2025): \u20AC0 \u2014 Development and pilot year. Invested in building the product and running 4 proof-of-concept events.',
    '',
    'Note: While we have not yet generated direct revenue, our pilot events represented significant economic value \u2014 GITS Cairo alone convened 170+ international trade leaders, and collectively our events facilitated B2B connections across 3 countries.',
])

# ── Q11 ──
add_question(11, 'Sales forecast N+1')
add_answer_block([
    'Year N+1 (2027) forecast: \u20AC15,000 \u2013 \u20AC25,000',
    '',
    'As a pre-revenue startup entering our first commercial year, this forecast is grounded in what we can realistically achieve:',
    '',
    '\u2022 2-3 Paid enterprise deployments for institutional partners from our existing network (GITS, IPDAYS, RedStart ecosystem events)',
    '\u2022 3-5 Small/mid-size event organizers onboarded through the Soft Landing program ecosystem on paid plans',
    '\u2022 Transaction commissions on ticketed events hosted on the platform',
    '\u2022 1 Pilot white-label or co-branded deployment with a European partner met through the program',
    '',
    "Pricing is being validated with early adopters during 2026. Our first commercial year will focus on securing reference clients and proving willingness-to-pay rather than scaling revenue. The real growth lever is the institutional network we already operate within \u2014 GITS events across 3 countries, RedStart's 800+ supported startups, and the Euro-African corridor partnerships built through this program.",
])

# ── Q12 ──
add_question(12, 'Capital stock')
add_fill_in('Capital stock amount and currency (e.g. EGP 100,000)')

# ── Q13 ──
add_question(13, 'Date of first sale of your product/service')
add_answer_block([
    'No commercial sales yet. First paid deployments targeted for Q3 2026, coinciding with European market entry through this program. Our pilot phase (May 2025 \u2013 present) focused on product validation and user acquisition.',
])

# ── Q14 ──
add_question(14, 'Number of customers and customer references (list a few names)')
add_answer_block([
    '1,200+ registered users across 4 pilot events. Current institutional partners and references:',
    '',
    '\u2022 GITS 2025 \u2014 Global Investment and Trade Summit (Cairo, Egypt & Abidjan, C\u00f4te d\'Ivoire) \u2014 250+ international leaders, organized with Egypt\'s Ministry of International Cooperation and GAFI Egypt',
    '\u2022 IPDAYS X GITS 2025 (Tunis, Tunisia) \u2014 Innovation & entrepreneurship summit organized with RedStart Tunisia at Radisson Hotel Tunis',
    '\u2022 RedStart Tunisia \u2014 Leading North African startup accelerator (800+ projects supported)',
    '\u2022 The LEE Experience (Lebanon) \u2014 Event and experience company',
    '\u2022 SEKETAK Solutions (Egypt) \u2014 Parent company and B2B consulting firm',
])

# ── Q15 ──
add_question(15, 'Have you been supported by an Incubator or an Accelerator?')
add_answer_block([
    'Yes. Eventra was developed within and supported by RedStart Tunisia, one of North Africa\'s most established startup accelerators, co-founded by our own Douja Gharbi. RedStart is recognized by the Tunisian government\'s Startup Tunisia initiative and has partnerships with UNIDO, INVESTMED, and the European ESIL program.',
    '',
    'Additionally, the platform was developed under SEKETAK Solutions (Egypt), which provides B2B consulting and entrepreneurship support services.',
])

# ── Q16 ──
add_question(16, 'Have you raised money for your startup? If yes, precise by whom?')
add_answer_block([
    'No external funding raised to date. The platform has been self-funded (bootstrapped) by the founding team. Development costs have been absorbed through SEKETAK Solutions\' operational budget and the founders\' personal investment.',
    '',
    'We are actively seeking our first funding round (pre-seed/seed) to finance European market entry, team expansion, and platform scaling. The Soft Landing program\'s investor network and fundraising acceleration are key motivations for our application.',
])

# ── Q17 ──
add_question(17, 'Have you obtained an Award or Label?')
add_answer_block([
    'While Eventra has not yet received formal awards, our pilot events have earned institutional recognition:',
    '',
    '\u2022 GITS 2025 was organized in partnership with Egypt\'s Ministry of International Cooperation and GAFI (General Authority for Investment and Free Zones), validating our platform\'s capacity to support government-level events.',
    '',
    '\u2022 IPDAYS X GITS 2025 received significant press coverage across Tunisian media (La Presse de Tunisie, Managers Magazine, Web Manager Center, Tekiano) as a landmark event connecting Tunisia\'s innovation ecosystem to international markets.',
])

# ── Q18 ──
add_question(18, 'Do you have an impact-driven vision?', required=False)
add_answer_block([
    'Absolutely. Eventra\'s core mission is to democratize professional event management for African entrepreneurs and institutions \u2014 removing the technology barrier that forces the continent\'s event industry to rely on expensive Western platforms not built for its realities.',
    '',
    'Our impact thesis:',
    '',
    '1. Economic Inclusion: By providing affordable, multilingual event infrastructure, we enable African SMEs, trade bodies, and institutions to run professional events that attract international attention and investment \u2014 leveling the playing field with European counterparts.',
    '',
    '2. B2B Bridge-Building: Every event on Eventra generates structured business connections between African and European actors. Our AI matchmaking doesn\'t just connect people \u2014 it creates measurable economic partnerships across the Mediterranean.',
    '',
    '3. Digital Sovereignty: We believe Africa\'s event data should be managed on platforms designed for Africa\'s needs. Eventra is built by an African team, with African languages, payment realities, and institutional frameworks in mind.',
    '',
    '4. Youth & Women Empowerment: Through our partnership with RedStart Tunisia (which has trained hundreds of women entrepreneurs), Eventra specifically supports events focused on youth entrepreneurship, women in business, and social innovation.',
    '',
    '5. Sustainable Events: Our all-digital approach reduces the paper, logistics, and travel waste associated with traditional event management \u2014 supporting the ecological transition in the events industry.',
])

# ── Q19 ──
add_question(19, 'If you have a demo, share the URL with us', required=False)
add_answer_block(['https://eventra.cloud/'])

# ── Q20 ──
add_question(20, 'Legal status')
add_fill_in('Legal status (e.g. Limited Liability Company / Egyptian law)')

# ── Q21 ──
add_question(21, 'Registering number')
add_fill_in('SEKETAK Solutions Egypt commercial registration number')

# ── Q22 ──
add_question(22, 'Headquarters address')
add_fill_in('SEKETAK Solutions registered address in Egypt')

# ── CHECKLIST ──
doc.add_page_break()
title2 = doc.add_heading('PRE-SUBMISSION CHECKLIST', level=1)
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER

checklist_items = [
    'Fill in % stakes for each founder (Q7)',
    'Fill in capital stock amount (Q12)',
    'Fill in legal status (Q20)',
    'Fill in registering number (Q21)',
    'Fill in headquarters address (Q22)',
    'Prepare slide deck \u2014 10 pages max, 30MB max PDF (Guy Kawasaki format)',
    'Review all answers for accuracy',
    'Submit on F6S before April 23, 2026',
]

for item in checklist_items:
    p = doc.add_paragraph()
    run = p.add_run(f'\u2610  {item}')
    run.font.size = Pt(12)

# ── SLIDE DECK OUTLINE ──
doc.add_paragraph()
doc.add_heading('RECOMMENDED SLIDE DECK STRUCTURE (10 pages)', level=1)

slides = [
    ('Slide 1 \u2014 Title', 'Eventra \u2014 The Event Platform Bridging Africa & Europe'),
    ('Slide 2 \u2014 Problem', 'African event organizers lack affordable, multilingual, all-in-one event tech. Western platforms don\'t fit African realities.'),
    ('Slide 3 \u2014 Solution', 'Eventra: no-code Design Studio + B2B matchmaking + trilingual (EN/FR/AR) + all-in-one management'),
    ('Slide 4 \u2014 Product Demo', 'Screenshots of Design Studio, landing pages, B2B matchmaking, dashboard'),
    ('Slide 5 \u2014 Business Model', 'Freemium SaaS + transaction fees + enterprise packages + B2B matchmaking premium'),
    ('Slide 6 \u2014 Traction', '4 events \u2022 3 countries \u2022 1,200+ users \u2022 Institutional partnerships (GITS, Ministry, GAFI)'),
    ('Slide 7 \u2014 Market', '\u20AC30B+ global event tech market, Africa fastest-growing segment, Marseille = Euro-African gateway'),
    ('Slide 8 \u2014 Team', 'Assem (CEO/B2B) + Douja (Ecosystem/RedStart) + Oussama (CTO/Product)'),
    ('Slide 9 \u2014 Competition', 'vs Eventbrite/Hopin/Bizzabo \u2014 our edge: trilingual, Africa-native, B2B matchmaking, affordable'),
    ('Slide 10 \u2014 Ask & Roadmap', '2026: EU setup + first revenue / 2027: scale + seed round / 2028: Series A'),
]

for title_text, content in slides:
    p = doc.add_paragraph()
    run = p.add_run(title_text + ': ')
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(content)
    run2.font.size = Pt(11)

# Save
output_path = os.path.join(
    os.environ.get('USERPROFILE', ''),
    'Desktop', 'Agent X', 'EVENTRA X', 'Eventra Code V1.1', 'eventra',
    'Soft_Landing_2026_Application_EVENTRA_v2.docx'
)
doc.save(output_path)
print(f'SUCCESS: Document saved to {output_path}')
