# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches

doc = Document()

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def sep():
    doc.add_paragraph('')

# TITLE
doc.add_heading('EVENTRA', 0)
add_para('Pitch Deck Content Guide', bold=True)
add_para('Soft Landing Provence Africa Connect 2026')
add_para('12-Slide Format  |  Max 30MB PDF')
add_para('This document contains the full text content, speaker notes, and visual direction for each slide.\nUse this to build the final presentation in PowerPoint, Google Slides, or Canva.', italic=True)

sep()

# ═══════════════════════════════════════════════════════════
# SLIDE 1 — TITLE
# ═══════════════════════════════════════════════════════════
doc.add_heading('SLIDE 1 \u2014 Title', 1)
doc.add_heading('EVENTRA \u2014 The All-in-One Event Management Platform', 2)
add_para('On-Screen Content', bold=True)
add_para('EVENTRA')
add_para('The All-in-One Event Management Platform')
add_para('Every event. One platform. Zero friction.', italic=True)
add_bullet('No-Code Website Builder  |  Ticketing  |  B2B Matchmaking  |  Check-In  |  Trilingual (EN/FR/AR)')
add_para('')
add_para('Soft Landing Provence Africa Connect 2026')
add_para('eventra.cloud')
add_para('\U0001f3a8 VISUAL:', bold=True)
add_para('Eventra logo centered on dark background (#0B2641). Blue accent (#0684F5). Abstract connected-nodes graphic suggesting global connectivity. Subtle globe or network motif. Bottom: SEKETAK logo. Clean, confident, modern.')
add_para('\U0001f399 SPEAKER NOTES:', bold=True)
add_para('Good morning. We are Eventra \u2014 the all-in-one event management platform. We cover the entire event lifecycle: from building your event website with zero code, to selling tickets, running email campaigns, powering B2B matchmaking on the day, and delivering analytics after. One platform replaces five disconnected tools. We are trilingual by architecture \u2014 English, French, and Arabic with full right-to-left support. The platform is live, battle-tested across 3 international summits in 4 countries, and we are here because Soft Landing is our gateway to Europe.')

sep()

# ═══════════════════════════════════════════════════════════
# SLIDE 2 — PROBLEM
# ═══════════════════════════════════════════════════════════
doc.add_heading('SLIDE 2 \u2014 The Problem', 1)
doc.add_heading('Event Management is Broken', 2)
add_para('Headline:', bold=True)
add_para('Running an event today means juggling 5+ disconnected tools \u2014 and still dropping the ball.')
add_para('')
add_para('The Story:', bold=True)
add_para('Imagine you are organizing a professional conference. You need to BUILD your event presence \u2014 so you hire a developer or fight with WordPress. You need TICKETING \u2014 so you set up Eventbrite. You need to COMMUNICATE \u2014 so you configure Mailchimp. On event day, you need CHECK-IN \u2014 another app. And B2B MATCHMAKING? That costs $18,000/year on enterprise platforms.')
add_para('')
add_para('Five tools. Five dashboards. Five invoices. Zero integration.', bold=True)
add_para('')
add_para('After the event? You copy data between spreadsheets to figure out what worked. This is how most event organizers operate today \u2014 whether in Paris, Cairo, Dubai, or Abidjan.')
add_para('')
add_para('The 5 Pain Points:', bold=True)
add_bullet('Building the event website \u2014 requires developers or expensive designers')
add_bullet('Ticketing & registration \u2014 separate platform, separate fees, no integration')
add_bullet('Communication \u2014 email tools completely disconnected from attendee data')
add_bullet('On-site check-in \u2014 manual processes or yet another app')
add_bullet('B2B matchmaking \u2014 $10K-$18K/year premium, or completely absent')
add_para('\U0001f3a8 VISUAL:', bold=True)
add_para('Left: frustrated organizer surrounded by floating tool logos (WordPress + Eventbrite + Mailchimp + Excel + WhatsApp + Brella) in chaotic cluster. Right: the 5 pain points as descending list with red X marks. Emotional, relatable, clean.')
add_para('\U0001f399 SPEAKER NOTES:', bold=True)
add_para('We lived this problem ourselves. When we organized GITS \u2014 the Global Investment and Trade Summit \u2014 across Cairo, Abidjan, and Tunis, we had to stitch together five or six tools every single time. Website on one platform, tickets on another, emails on a third, check-in on a fourth, and for B2B matchmaking? The options started at 10,000 dollars a year. We spent more time managing tools than managing the actual event. That frustration is what led us to build Eventra.')

sep()

# ═══════════════════════════════════════════════════════════
# SLIDE 3 — MARKET OPPORTUNITY
# ═══════════════════════════════════════════════════════════
doc.add_heading('SLIDE 3 \u2014 Market Opportunity', 1)
doc.add_heading('A $9.3B Market with a Massive Underserved Middle', 2)
add_para('TAM \u2014 Total Addressable Market:', bold=True)
add_bullet('Global Event Management Software: \u20ac8.6B in 2025 \u2192 \u20ac16B by 2030')
add_bullet('13.2% CAGR (Source: Grand View Research, 2025)')
add_para('')
add_para('SAM \u2014 Serviceable Addressable Market:', bold=True)
add_bullet('Europe + MENA + Africa mid-market segment: \u20ac2.1B')
add_bullet('Organizers running 2-50 events/year, priced out of enterprise tools')
add_para('')
add_para('SOM \u2014 Serviceable Obtainable Market (3-year target):', bold=True)
add_bullet('120 paying clients \u00d7 avg \u20ac1,400/year = \u20ac168K ARR by Year 3')
add_bullet('Entry via Africa-Europe corridor, expand globally')
add_para('')
add_para('European Market:', bold=True)
add_bullet('European EMS Market: \u20ac3.1B in 2025, growing at 14.6% CAGR')
add_bullet('France: 2nd largest event market in Europe after Germany')
add_para('')
add_para('Key Growth Drivers:', bold=True)
add_bullet('Hybrid event adoption (60%+ of conferences offer both formats)')
add_bullet('AI integration (45% of new platforms embed AI features)')
add_bullet('SME demand for affordable all-in-one solutions')
add_bullet('Post-pandemic "experience economy" driving record event volumes')
add_para('')
add_para('The Gap:', bold=True)
add_para('Enterprise tools (Cvent, Bizzabo) serve the top 5% at \u20ac10K-20K/year. Basic tools (Eventbrite) serve simple ticketing. The 80% in between \u2014 professional organizers who need a full platform but cannot justify enterprise pricing \u2014 are completely underserved. That is Eventra\u2019s market.')
add_para('\U0001f3a8 VISUAL:', bold=True)
add_para('TAM/SAM/SOM funnel diagram (large to small). Three stat boxes (\u20ac8.6B / \u20ac2.1B / \u20ac168K target). Growth chart 2025-2030. Clean, data-driven.')
add_para('\U0001f399 SPEAKER NOTES:', bold=True)
add_para('The global market is 8.6 billion euros today, growing to 16 billion by 2030. But let me focus on where we actually compete. Our serviceable market is the Europe, MENA, and Africa mid-market \u2014 organizers priced out of 20,000-euro enterprise tools but who need more than basic ticketing. That segment alone is 2.1 billion euros. Our realistic 3-year target is 120 paying clients. The opportunity is clear: enterprise tools serve the top 5 percent, basic tools serve simple events, and 80 percent of professional organizers are left with nothing. We fill that gap.')

sep()

# ═══════════════════════════════════════════════════════════
# SLIDE 4 — COMPETITIVE LANDSCAPE
# ═══════════════════════════════════════════════════════════
doc.add_heading('SLIDE 4 \u2014 Competitive Landscape', 1)
doc.add_heading('They Built for Enterprises. We Built for Everyone.', 2)
add_para('POSITIONING MAP (2\u00d72 Matrix):', bold=True)
add_para('X-axis: Price (Low \u2192 High) | Y-axis: Feature Completeness (Low \u2192 High)')
add_para('')
add_bullet('Bottom-Left (Low price, Low features): Eventbrite, Ticket Tailor')
add_bullet('Top-Right (High price, High features): Cvent, Bizzabo, Hubilo')
add_bullet('Top-Left (Low price, High features): EVENTRA \u2190 This is the white space')
add_bullet('Bottom-Right (High price, Low features): Hopin/RingCentral (virtual-only)')
add_para('')
add_para('TOP 5 COMPETITOR COMPARISON (simplified for slide):', bold=True)

table = doc.add_table(rows=7, cols=7)
table.style = 'Table Grid'
headers = ['Platform', 'Price/Year', 'No-Code Builder', 'Multilingual', 'AI Matchmaking', 'All-in-One']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ['Eventbrite', '\u20ac0 + 3.7%/tkt', 'Basic templates', '7 langs (no AR)', 'No', 'No (ticketing only)'],
    ['Bizzabo', '\u20ac16,500+/yr', 'Yes (strong)', '500+ langs', 'Yes (AI)', 'Yes'],
    ['Cvent', '\u20ac18,000+/yr', 'Yes (best)', 'AI translate', 'Yes', 'Yes'],
    ['Swapcard', '\u20ac1,300+/yr', 'Partial', 'EN/FR/ES', 'Yes (AI)', 'Yes'],
    ['Hubilo', '\u20ac9,200+/yr', 'Yes', 'Limited', 'Yes (AI)', 'Yes'],
    ['EVENTRA', 'FREE \u2013 \u20ac540/yr', 'Yes (15+ blocks)', 'EN/FR/AR + RTL', 'Yes (AI)', 'Yes'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, cell_data in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = cell_data

add_para('')
add_para('Our Unfair Advantages:', bold=True)
add_bullet('Only platform with native Arabic + full RTL support (400M+ Arabic speakers globally)')
add_bullet('Enterprise-grade features at 30-38x lower price than Bizzabo/Cvent')
add_bullet('Purpose-built for trade missions and institutional B2B events')
add_bullet('Cost structure advantage: lean team, no VC bloat, emerging-market efficiency')
add_para('')
add_para('Full 10-competitor analysis available in appendix.', italic=True)
add_para('\U0001f3a8 VISUAL:', bold=True)
add_para('2\u00d72 positioning matrix (main visual) with competitor logos plotted. Eventra in the "white space" top-left quadrant highlighted in blue. Simplified comparison table below. Clean, scannable.')
add_para('\U0001f399 SPEAKER NOTES:', bold=True)
add_para('Here is where Eventra sits in the competitive landscape. The market has two extremes: cheap but limited tools like Eventbrite \u2014 just ticketing, no builder, no matchmaking \u2014 and full-featured enterprise platforms like Cvent and Bizzabo at 16 to 18 thousand euros per year. The white space is top-left: full features at accessible pricing. That is where we are. And we have an additional edge no competitor has: native Arabic with full right-to-left support. That opens a 400-million-person market that every single competitor ignores. Our cost advantage is structural: lean team, no investor bloat, and emerging-market development efficiency.')

sep()

# ═══════════════════════════════════════════════════════════
# SLIDE 5 — SOLUTION (3 PHASES)
# ═══════════════════════════════════════════════════════════
doc.add_heading('SLIDE 5 \u2014 Our Solution: The 3-Phase Event Engine', 1)
doc.add_heading('One Platform. Entire Lifecycle. Zero Code.', 2)
add_para('Headline:', bold=True)
add_para('Eventra replaces 5 tools with one \u2014 covering Pre-Event, During Event, and Post-Event.')
add_para('')
add_para('\U0001f680 PRE-EVENT \u2014 Build & Promote', bold=True)
add_bullet('No-Code Website Builder \u2014 15+ drag-and-drop blocks (hero, speakers, agenda, sponsors, tickets, FAQ, countdown...)')
add_bullet('Registration & Ticketing \u2014 Free/paid/VIP/early-bird, promo codes, custom forms')
add_bullet('Marketing Emails \u2014 Enterprise-grade campaigns with templates and scheduling')
add_bullet('Sponsorship & Speaker Management \u2014 Tiers, profiles, multi-track agendas')
add_bullet('B2B Matchmaking Setup \u2014 Attendee profiling, interest matching, meeting pre-scheduling')
add_para('')
add_para('\u26a1 DURING EVENT \u2014 Manage & Connect', bold=True)
add_bullet('On-Site QR Check-In \u2014 Real-time scanning and attendance tracking')
add_bullet('Live B2B Matchmaking \u2014 AI-powered meeting suggestions, 1:1 scheduling')
add_bullet('Trade Mission Tools \u2014 Delegation management, bilateral coordination')
add_bullet('Real-Time Dashboard \u2014 Live metrics, check-in rates, engagement')
add_para('')
add_para('\U0001f4ca POST-EVENT \u2014 Analyze & Follow-up', bold=True)
add_bullet('Analytics & ROI Reporting \u2014 Attendance funnels, revenue, sponsor ROI')
add_bullet('Data Export \u2014 CSV/Excel, CRM-ready')
add_bullet('Follow-up Campaigns \u2014 Thank-you emails, surveys, next-event promotion')
add_bullet('One-Click Duplication \u2014 Clone events for recurring series')
add_para('\U0001f3a8 VISUAL:', bold=True)
add_para('Three-column layout (Pre \u2192 During \u2192 Post) with connecting arrow flow. 3-4 features per column with icons. Actual Eventra UI screenshots embedded subtly behind each phase. Clean, not crowded.')
add_para('\U0001f399 SPEAKER NOTES:', bold=True)
add_para('Eventra is not a website builder or a ticketing tool \u2014 it is a complete event engine. Before the event: build your entire online presence with 15-plus drag-and-drop blocks, set up ticketing, launch email campaigns, configure matchmaking. During: QR check-in, live AI-powered B2B meetings, trade mission coordination, real-time dashboard. After: full analytics, data export, follow-up emails, and one-click event duplication for your next edition. Everything that used to require 5 separate tools and 5 separate subscriptions now lives in one platform, accessible in 3 languages.')

sep()

# ═══════════════════════════════════════════════════════════
# SLIDE 6 — PRODUCT DEMO
# ═══════════════════════════════════════════════════════════
doc.add_heading('SLIDE 6 \u2014 Product in Action', 1)
doc.add_heading('Live Product. Real Events. Not Vaporware.', 2)
add_para('Headline:', bold=True)
add_para('Not a mockup. Not a prototype. This is live at eventra.cloud.')
add_para('')
add_para('4 Key Screenshots (full-bleed in device mockups):', bold=True)
add_para('')
add_para('1. NO-CODE WEBSITE BUILDER', bold=True)
add_para('Drag-and-drop event website builder with 15+ blocks. Full customization, instant preview, mobile-responsive output.')
add_para('')
add_para('2. PUBLISHED EVENT LANDING PAGE', bold=True)
add_para('A real event page built by our platform \u2014 responsive, trilingual, with integrated ticketing and registration.')
add_para('')
add_para('3. B2B MATCHMAKING ENGINE', bold=True)
add_para('AI-powered attendee matching based on profiles, sectors, and business goals. Meeting scheduling built in.')
add_para('')
add_para('4. ON-SITE CHECK-IN SCANNER', bold=True)
add_para('QR code scanning interface with real-time attendance tracking and dashboard.')
add_para('')
add_para('Live Demo: eventra.cloud', bold=True)
add_para('\U0001f3a8 VISUAL:', bold=True)
add_para('4 screenshots in a 2\u00d72 grid, each in a device mockup (laptop or phone). Real UI, not wireframes. The URL "eventra.cloud" prominent at bottom. Goal: prove this is REAL and POLISHED.')
add_para('\U0001f399 SPEAKER NOTES:', bold=True)
add_para('Let me show you the product. Top-left: our no-code website builder where organizers drag and drop blocks to build their event website in minutes. Top-right: a real published event page \u2014 this is what attendees see, fully responsive, with integrated registration and ticketing. Bottom-left: our B2B matchmaking engine that connects the right people based on their profiles. Bottom-right: our on-site check-in scanner that gives organizers a live dashboard. Everything you see is live at eventra.cloud and has been used at real summits.')

sep()

# ═══════════════════════════════════════════════════════════
# SLIDE 7 — TRACTION
# ═══════════════════════════════════════════════════════════
doc.add_heading('SLIDE 7 \u2014 Traction & Proof', 1)
doc.add_heading('Battle-Tested Across 3 International Summits', 2)
add_para('Headline:', bold=True)
add_para('We did not come with a pitch. We came with proof.')
add_para('')
add_para('Key Metrics:', bold=True)
add_bullet('3 major international summits powered by Eventra')
add_bullet('4 countries: Egypt, C\u00f4te d\'Ivoire, Tunisia, France')
add_bullet('670+ attendees (government officials, CEOs, investors, entrepreneurs)')
add_bullet('60+ structured B2B meetings facilitated')
add_bullet('6+ national press outlets covered our events')
add_bullet('Institutional partners: Egypt Ministry, GAFI, AFD, CEPEX, AfriLabs')
add_para('')

add_para('EVENT 1: GITS 2025 \u2014 Cairo, Egypt', bold=True)
add_bullet('Global Investment & Trade Summit | May 12-13, 2025 | Marriott Zamalek')
add_bullet('170 global trade leaders | 6 countries | 60 B2B meetings')
add_bullet('Partners: Egypt Ministry of International Cooperation, GAFI Egypt')
add_bullet('Sectors: Fintech, Agritech, Solar, Food Processing, E-commerce')
add_para('')

add_para('EVENT 2: GITS 2025 \u2014 Abidjan, C\u00f4te d\'Ivoire', bold=True)
add_bullet('Global Investment & Trade Summit \u2014 West Africa Edition')
add_bullet('250+ changemakers and decision-makers')
add_bullet('Focus: Cross-border trade, sustainable growth, women entrepreneurship')
add_bullet('Partners: SEKETAK, RedStart Tunisia, AfriLabs, JEUN\'ESS Market Fund')
add_para('')

add_para('EVENT 3: IPDAYS x GITS 2025 \u2014 Tunis, Tunisia', bold=True)
add_bullet('Innovation & Partnership Days (4th edition) | Nov 12-13, 2025 | Radisson Hotel')
add_bullet('250+ entrepreneurs, investors, and institutional partners')
add_bullet('Key speakers: Minister of Commerce, AFD, CEPEX, Konza Technopolis (Kenya)')
add_bullet('Partners: RedStart Tunisia, Expertise France, AFD')
add_bullet('Press: La Presse, Managers, WMC, Tekiano, R\u00e9alit\u00e9s, Kapitalis')
add_para('')

add_para('STATUS: Pilot phase complete. Commercial launch Q3 2026.', bold=True)
add_para('These events validated product-market fit at scale. We are now transitioning from free pilots to paid commercial deployments, with existing relationships forming our initial sales pipeline.')
add_para('\U0001f3a8 VISUAL:', bold=True)
add_para('Top: 4-5 key metric counters (large numbers). Middle: 3 event cards side by side, each with flag, city, attendee count, partner logos. Bottom: "Pilot Complete \u2192 Commercial Launch Q3 2026" timeline arrow. World map with pins.')
add_para('\U0001f399 SPEAKER NOTES:', bold=True)
add_para('We are not pitching an idea. In 2025, Eventra powered three major international summits. In Cairo: 170 trade leaders from 6 countries with 60 B2B meetings, organized with Egypt\u2019s Ministry of International Cooperation. In Abidjan: 250 participants focused on West African cross-border trade. In Tunis: 250 participants with the Minister of Commerce, the French Development Agency, and coverage from six national media outlets. These were our pilot events \u2014 deployed free to validate the platform at institutional scale. The product works. Now we are transitioning to paid commercial deployments starting Q3 2026, and our existing network of organizers from GITS and RedStart forms our initial pipeline.')

sep()

# ═══════════════════════════════════════════════════════════
# SLIDE 8 — BUSINESS MODEL & PRICING
# ═══════════════════════════════════════════════════════════
doc.add_heading('SLIDE 8 \u2014 Business Model & Pricing', 1)
doc.add_heading('Enterprise Features. Accessible Pricing.', 2)
add_para('Revenue Model:', bold=True)
add_bullet('70% Subscriptions (recurring SaaS revenue)')
add_bullet('25% Transaction fees (2-5% per paid ticket)')
add_bullet('5% Premium add-ons (white-label, API, dedicated support)')
add_para('')
add_para('THREE PLANS:', bold=True)
add_para('')
add_para('STARTER \u2014 Free', bold=True)
add_para('1 event, 3-day max, basic features. Conversion funnel entry point.')
add_para('')
add_para('PRO \u2014 \u20ac45/month', bold=True)
add_para('5 events/mo, full features, unlimited emails, B2B matchmaking, advanced analytics.')
add_bullet('6-month commitment: \u20ac40/mo (10% off)')
add_bullet('Annual commitment: \u20ac36/mo (20% off = \u20ac432/year)')
add_para('')
add_para('ENTERPRISE \u2014 \u20ac275/month', bold=True)
add_para('Unlimited events, unlimited duration, trade mission tools, white-label, dedicated support.')
add_bullet('6-month commitment: \u20ac248/mo (10% off)')
add_bullet('Annual commitment: \u20ac220/mo (20% off = \u20ac2,640/year)')
add_para('')
add_para('Conversion Strategy:', bold=True)
add_para('Free \u2192 Pro: Users hit event/duration limits on their second event. Natural upgrade trigger.')
add_para('Pro \u2192 Enterprise: Agencies and institutions running 5+ events/month need unlimited scale.')
add_para('')
add_para('Why We Can Be 30x Cheaper:', bold=True)
add_bullet('Lean 3-person founding team (no VC-funded headcount bloat)')
add_bullet('Modern serverless architecture (Supabase) = minimal infrastructure costs')
add_bullet('Emerging-market development efficiency')
add_bullet('No legacy systems to maintain')
add_para('')
add_para('Price Comparison:', bold=True)
add_bullet('Eventra Pro (\u20ac432/yr) vs Bizzabo (\u20ac16,500/yr) = 38x cheaper')
add_bullet('Eventra Enterprise (\u20ac2,640/yr) vs Cvent (\u20ac18,000/yr) = 7x cheaper')
add_para('\U0001f3a8 VISUAL:', bold=True)
add_para('Three pricing cards (Starter/Pro/Enterprise) with key features listed. Pro highlighted as "Most Popular." Below: conversion funnel diagram (Free \u2192 Pro \u2192 Enterprise). Price comparison callout vs competitors.')
add_para('\U0001f399 SPEAKER NOTES:', bold=True)
add_para('Our business model is straightforward SaaS. Three plans with flexible billing. Starter is free \u2014 it is our funnel entry point. Once an organizer runs their second event and hits the limits, they naturally upgrade to Pro at 45 euros per month \u2014 or 36 on an annual plan. Enterprise at 275 per month serves agencies and institutions running events continuously. A common question: how can you be 38 times cheaper than Bizzabo while offering comparable features? Three reasons: we are a lean founding team with zero bloat, we use modern serverless infrastructure that costs a fraction of legacy systems, and we developed in an emerging market with structural cost advantages. This is not unsustainable pricing \u2014 it is efficient pricing.')

sep()

# ═══════════════════════════════════════════════════════════
# SLIDE 9 — GO-TO-MARKET STRATEGY
# ═══════════════════════════════════════════════════════════
doc.add_heading('SLIDE 9 \u2014 Go-to-Market Strategy', 1)
doc.add_heading('How We Get to 120 Clients in 3 Years', 2)
add_para('Strategy: Global product. Africa-Europe beachhead.', bold=True)
add_para('')
add_para('PHASE 1: Activate Existing Network (H2 2026)', bold=True)
add_bullet('Convert GITS summit partners into paying clients (direct relationships with 20+ institutional organizers)')
add_bullet('Leverage RedStart Tunisia network (800+ supported startups = event organizers)')
add_bullet('Soft Landing introductions to French chambers of commerce and event organizers')
add_bullet('Target: 10 paying clients from warm pipeline')
add_para('')
add_para('PHASE 2: Expand via Partnerships (2027)', bold=True)
add_bullet('Partner with accelerators and incubators (who organize 10-20 events/year each)')
add_bullet('Chamber of commerce partnerships in France and MENA')
add_bullet('Content marketing + SEO in French and Arabic (zero competition in these languages)')
add_bullet('Showcase at Le Grand Bain + Emerging Valley (via Soft Landing)')
add_bullet('Target: 40 paying clients')
add_para('')
add_para('PHASE 3: Product-Led Growth (2028+)', bold=True)
add_bullet('Free tier drives organic signups globally')
add_bullet('Arabic/French SEO moat (no competitor targets these)')
add_bullet('Referral program: organizers recommend to other organizers')
add_bullet('API partnerships with CRM and marketing platforms')
add_bullet('Target: 120+ paying clients')
add_para('')
add_para('Why This Works:', bold=True)
add_bullet('We already HAVE the relationships \u2014 GITS network spans 4 countries')
add_bullet('RedStart\u2019s 800 startups need event tools = built-in demand')
add_bullet('Zero competition in French/Arabic event management SEO')
add_bullet('Soft Landing provides the European launchpad we need')
add_para('\U0001f3a8 VISUAL:', bold=True)
add_para('3-phase horizontal timeline (Activate \u2192 Expand \u2192 Scale). Each phase shows channel icons and client targets. Map showing geographic expansion. Partnership logos.')
add_para('\U0001f399 SPEAKER NOTES:', bold=True)
add_para('Let me explain exactly how we get customers. Phase 1 is the lowest-hanging fruit: we already have direct relationships with 20-plus institutional organizers from our GITS summits. These are ministries, chambers of commerce, trade bodies who used our platform for free during pilots. Converting them to paid is our first motion. Add RedStart\u2019s network of 800 startups \u2014 all of whom organize events and need tools. Plus Soft Landing introductions to French partners. Phase 2: we partner with accelerators and incubators who each run 10-20 events per year. We become their default event tool. Phase 3: product-led growth. Our free tier drives signups, our French and Arabic SEO gives us a moat no competitor has, and referrals compound. The path from 0 to 120 clients is not a hope \u2014 it is a plan built on existing assets.')

sep()

# ═══════════════════════════════════════════════════════════
# SLIDE 10 — TEAM
# ═══════════════════════════════════════════════════════════
doc.add_heading('SLIDE 10 \u2014 The Team', 1)
doc.add_heading('Built by Operators Who Live This Problem', 2)
add_para('FOUNDING TEAM:', bold=True)
add_para('')

add_para('Assem Kamel \u2014 CEO & Co-Founder', bold=True)
add_bullet('Serial entrepreneur: Co-founder of GITS and The LEE Experience (Lebanon)')
add_bullet('Certified Business Coach (ILO \u2014 International Labour Organization)')
add_bullet('15+ years in international trade and cross-border business development')
add_bullet('Network: Ministries, GAFI Egypt, chambers of commerce across MENA and Africa')
add_para('')

add_para('Douja Gharbi \u2014 Co-Founder & Strategic Partnerships', bold=True)
add_bullet('CEO of RedStart Tunisia \u2014 North Africa\u2019s leading startup accelerator')
add_bullet('30+ years in entrepreneurship and economic development')
add_bullet('800+ projects supported, ESIL Angel for Tunisia')
add_bullet('Partnerships: UNIDO, INVESTMED, AFD, Expertise France')
add_para('')

add_para('Oussama Lamine \u2014 CTO', bold=True)
add_bullet('Full-stack developer and product architect (React/TypeScript, Supabase, AI)')
add_bullet('Built Eventra end-to-end: no-code website builder, B2B matchmaking, email system')
add_bullet('5+ years in software development, specialized in SaaS platforms')
add_bullet('Sole technical lead delivering production-grade platform across 3 international events')
add_para('')

add_para('Why This Team Wins:', bold=True)
add_bullet('We organized GITS ourselves across 3 countries \u2014 we are our own first customer')
add_bullet('50+ combined years in events, trade, technology, and ecosystem building')
add_bullet('Existing network: 800+ startups, government ministries, European institutions')
add_bullet('Complementary: Business (Assem) + Ecosystem (Douja) + Technology (Oussama)')
add_para('')
add_para('Company: SEKETAK Solutions | HQ: Egypt', bold=True)
add_para('Seeking: European entity via Soft Landing Provence Africa Connect')
add_para('')
add_para('Key hire planned (post-program): European Commercial Lead based in Aix-Marseille-Provence for client acquisition in France and EU markets.', italic=True)
add_para('\U0001f3a8 VISUAL:', bold=True)
add_para('Professional headshots in a row with name, title, 2 key credentials below each. "Why This Team Wins" as sidebar bullets. Logos: SEKETAK, RedStart, GITS. Clean, personal, confident.')
add_para('\U0001f399 SPEAKER NOTES:', bold=True)
add_para('Our team has three complementary strengths. Assem brings 15 years of international trade and B2B relationships \u2014 he co-founded GITS and has direct access to ministries and investment authorities across four countries. Douja brings 30 years of ecosystem building as CEO of RedStart Tunisia \u2014 800 startups supported, partnerships with UNIDO and the French Development Agency, and deep European institutional connections. Oussama built the entire platform from scratch \u2014 frontend, backend, AI matchmaking, email system \u2014 and delivered it at production quality across three international summits. We are not outsiders theorizing about events. We ran GITS. We lived the pain. We built the solution. And through Soft Landing, we are hiring our first European commercial lead in Aix-Marseille-Provence.')

sep()

# ═══════════════════════════════════════════════════════════
# SLIDE 11 — FINANCIAL ROADMAP
# ═══════════════════════════════════════════════════════════
doc.add_heading('SLIDE 11 \u2014 Financial Roadmap', 1)
doc.add_heading('Conservative Projections. Clear Path.', 2)
add_para('All figures in EUR. Based on annual billing (20% discount applied).', italic=True)
add_para('')

add_para('Year 1 (H2 2026 \u2013 H1 2027): ESTABLISH', bold=True)
add_bullet('Clients: 10 (6 Pro + 4 Enterprise)')
add_bullet('Subscription revenue: 6 \u00d7 \u20ac432 + 4 \u00d7 \u20ac2,640 = \u20ac13,152')
add_bullet('Transaction fees (\u20ac150K ticket volume \u00d7 3%): \u20ac4,500')
add_bullet('Total Year 1 revenue: ~\u20ac18,000')
add_bullet('Milestone: First European paying clients via Soft Landing network')
add_para('')

add_para('Year 2 (2027\u20132028): GROW', bold=True)
add_bullet('Clients: 40 (25 Pro + 15 Enterprise)')
add_bullet('Subscription revenue: 25 \u00d7 \u20ac432 + 15 \u00d7 \u20ac2,640 = \u20ac50,400')
add_bullet('Transaction fees (\u20ac600K ticket volume \u00d7 3%): \u20ac18,000')
add_bullet('Total Year 2 revenue: ~\u20ac68,000')
add_bullet('Hire: European Commercial Lead + Customer Success')
add_bullet('Milestone: Pre-seed round (\u20ac200\u2013400K)')
add_para('')

add_para('Year 3 (2028\u20132029): SCALE', bold=True)
add_bullet('Clients: 120 (75 Pro + 45 Enterprise)')
add_bullet('Subscription revenue: 75 \u00d7 \u20ac432 + 45 \u00d7 \u20ac2,640 = \u20ac151,200')
add_bullet('Transaction fees (\u20ac2M ticket volume \u00d7 3%): \u20ac60,000')
add_bullet('Total Year 3 revenue: ~\u20ac211,000')
add_bullet('Team: 6\u20138 people')
add_bullet('Milestone: Seed round (\u20ac500K\u20131M)')
add_bullet('Expansion: Portuguese, Spanish, Turkish')
add_para('')

add_para('Runway & Burn Rate:', bold=True)
add_bullet('Current monthly burn: ~\u20ac2,000 (infrastructure + tools, team is equity-compensated)')
add_bullet('Pre-revenue runway: 12+ months with current bootstrapped model')
add_bullet('Post pre-seed: 18\u201324 months runway to reach profitability')
add_para('')

add_para('Path to Profitability:', bold=True)
add_para('Break-even at ~30 clients (mix of Pro + Enterprise). Projected: Month 14\u201318 post-commercial launch.')
add_para('')

add_para('Unit Economics:', bold=True)
add_bullet('Customer Acquisition Cost (target): \u20ac150\u2013300 (B2B institutional sales cycle)')
add_bullet('Avg LTV (blended Pro + Enterprise, 2-year retention): \u20ac2,800')
add_bullet('LTV:CAC ratio: 9\u201319x')
add_bullet('Gross margin: 85%+ (SaaS, serverless infrastructure)')
add_para('\U0001f3a8 VISUAL:', bold=True)
add_para('Revenue bar chart (Y1: \u20ac18K \u2192 Y2: \u20ac68K \u2192 Y3: \u20ac211K) stacked by subscriptions + transaction fees. Break-even line marked. Runway indicator. Clean financial slide.')
add_para('\U0001f399 SPEAKER NOTES:', bold=True)
add_para('Let me walk through the numbers. Year 1: 10 clients from our existing network \u2014 mostly GITS partners and RedStart contacts. That gives us 18,000 euros. Conservative. Year 2: we expand to 40 clients through partnerships and Soft Landing introductions \u2014 68,000 euros and we raise a pre-seed round. Year 3: product-led growth kicks in, we reach 120 clients and 211,000 euros, and we raise a seed round to accelerate. Important context: our current burn rate is only 2,000 euros per month because the team is equity-compensated during this phase. We break even at approximately 30 clients, which we expect to hit within 14 to 18 months of commercial launch. Our CAC is realistic for B2B institutional sales \u2014 150 to 300 euros \u2014 and our blended lifetime value of 2,800 euros gives us a healthy 9 to 19x return on acquisition spend.')

sep()

# ═══════════════════════════════════════════════════════════
# SLIDE 12 — THE ASK
# ═══════════════════════════════════════════════════════════
doc.add_heading('SLIDE 12 \u2014 The Ask', 1)
doc.add_heading('From Soft Landing to European Market Leader', 2)
add_para('WHAT WE NEED FROM SOFT LANDING:', bold=True)
add_bullet('Legal entity setup in Aix-Marseille-Provence for European operations')
add_bullet('B2B introductions to event organizers, chambers of commerce, and institutional partners in France')
add_bullet('Investor network access for our pre-seed round')
add_bullet('Mentorship on European market entry, GDPR compliance, and commercial strategy')
add_bullet('Stage at Le Grand Bain and Emerging Valley to showcase Eventra')
add_para('')
add_para('WHAT WE BRING:', bold=True)
add_bullet('A live, production-grade platform \u2014 not a prototype')
add_bullet('3 international summits delivered across 4 countries')
add_bullet('Institutional credibility: government ministries, GAFI, AFD, CEPEX')
add_bullet('A team with 50+ years combined experience and a network spanning Africa, MENA, and Europe')
add_bullet('A clear market gap validated by 670+ users and 6+ press outlets')
add_para('')
add_para('GDPR READINESS:', bold=True)
add_para('Platform architecture supports EU data compliance. European entity enables full GDPR-compliant data processing for EU clients. Implementation timeline: Q3-Q4 2026.')
add_para('')
add_para('PIPELINE:', bold=True)
add_para('3 LOI discussions in progress with institutional organizers from our GITS network for Q3-Q4 2026 paid deployments.')
add_para('')
add_para('CALL TO ACTION:', bold=True)
add_para('')
add_para('We are not asking you to bet on an idea.', bold=True)
add_para('We are asking you to back a product that already works,')
add_para('a team that already delivers,')
add_para('and a market that is already growing.')
add_para('')
add_para('eventra.cloud  |  assem@seketak-eg.com', bold=True)
add_para('\U0001f3a8 VISUAL:', bold=True)
add_para('Left: "What We Need" as a clean checklist with checkboxes. Right: "What We Bring" as proof points with green checkmarks. Bottom: bold call-to-action text, Eventra logo, contact info. Final impression: confidence, clarity, readiness.')
add_para('\U0001f399 SPEAKER NOTES:', bold=True)
add_para('Here is our ask. Five things. Help us set up our European entity. Introduce us to French institutional event organizers. Connect us with your investor network. Mentor us on European market entry and GDPR. And give us a stage at Le Grand Bain and Emerging Valley. In return, we bring something most applicants cannot: a live platform that has already powered international summits with government-level partners. We are GDPR-ready architecturally and will complete compliance by Q4 2026. We have three letters of intent in discussion with organizers from our existing network for paid deployments later this year. We are not asking you to take a risk. We are asking you to accelerate something that is already working. Thank you.')

sep()

# ═══════════════════════════════════════════════════════════
# APPENDIX
# ═══════════════════════════════════════════════════════════
doc.add_heading('APPENDIX', 1)
add_para('')
doc.add_heading('A. Full 10-Competitor Comparison Table', 2)

full_table = doc.add_table(rows=12, cols=9)
full_table.style = 'Table Grid'
headers = ['Platform', 'Price/Year', 'No-Code Builder', 'Multilingual', 'AI Match', 'Ticketing', 'Email Mktg', 'Check-in', 'All-in-One']
for i, h in enumerate(headers):
    full_table.rows[0].cells[i].text = h

data = [
    ['Eventbrite', '3.7%+\u20ac1.65/tkt', 'Basic templates', '7 langs (no AR)', 'No', 'Yes', 'Basic', 'Yes', 'No'],
    ['Hopin/RingCentral', '\u20ac1,100+', 'Basic', 'No', 'No', 'Yes', 'Yes', 'No', 'Virtual only'],
    ['Bizzabo', '\u20ac16,500+', 'Yes (strong)', '500+ langs', 'Yes (AI)', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['Cvent', '\u20ac18,000+', 'Yes (best)', 'AI translate', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['Splash', 'Custom', 'Yes (design)', 'Partial', 'No', 'Yes', 'Yes', 'Yes', 'Partial'],
    ['Whova', 'Custom/event', 'Basic', 'Limited', 'Partial', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['Airmeet', '\u20ac980+', 'Partial', 'Subtitles', 'Yes (AI)', 'Yes', 'Basic', 'No', 'Virtual'],
    ['Swapcard', '\u20ac1,300+', 'Partial', 'EN/FR/ES', 'Yes (AI)', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['Hubilo', '\u20ac9,200+', 'Yes', 'Limited', 'Yes (AI)', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['Accelevents', '\u20ac460+/event', 'Yes', 'Limited', 'Partial', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['EVENTRA', 'FREE\u2013\u20ac432/yr', '15+ blocks', 'EN/FR/AR+RTL', 'Yes (AI)', 'Yes', 'Yes', 'Yes', 'YES'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, cell_data in enumerate(row_data):
        full_table.rows[row_idx].cells[col_idx].text = cell_data

add_para('')
doc.add_heading('B. Full Pricing Feature Table', 2)

ptable = doc.add_table(rows=19, cols=4)
ptable.style = 'Table Grid'
pheaders = ['Feature', 'STARTER (Free)', 'PRO (\u20ac45/mo)', 'ENTERPRISE (\u20ac275/mo)']
for i, h in enumerate(pheaders):
    ptable.rows[0].cells[i].text = h

pricing_data = [
    ['No-Code Website Builder', '\u2713', '\u2713', '\u2713'],
    ['Number of Events', '1 event', '5 events/mo', 'Unlimited'],
    ['Event Duration', '3 days max', '14 days max', 'Unlimited'],
    ['Registration & Ticketing', '\u2713', '\u2713', '\u2713'],
    ['Custom Registration Forms', 'Basic', 'Advanced', 'Advanced'],
    ['Email Campaigns', '1 campaign', 'Unlimited', 'Unlimited'],
    ['Speaker Management', '5 speakers', 'Unlimited', 'Unlimited'],
    ['Session/Agenda Management', '\u2713', '\u2713', '\u2713'],
    ['Sponsorship Management', '\u2717', '\u2713', '\u2713'],
    ['B2B Matchmaking', '\u2717', '\u2713', '\u2713 (AI-enhanced)'],
    ['Trade Mission Tools', '\u2717', '\u2717', '\u2713'],
    ['On-Site QR Check-In', '\u2713', '\u2713', '\u2713'],
    ['Enterprise Email System', '\u2717', '\u2713', '\u2713'],
    ['Analytics & Reporting', 'Basic', 'Advanced', 'Advanced + Export'],
    ['Attendee Data Export', '\u2717', '\u2713', '\u2713'],
    ['White-label Branding', '\u2717', '\u2717', '\u2713'],
    ['Priority Support', '\u2717', '\u2713', '\u2713 (Dedicated)'],
    ['Payment Processing Fee', '5%', '3%', '2% (fixed)'],
]
for row_idx, row_data in enumerate(pricing_data, 1):
    for col_idx, cell_data in enumerate(row_data):
        ptable.rows[row_idx].cells[col_idx].text = cell_data

add_para('')
doc.add_heading('C. Design Notes', 2)
add_para('Brand Colors:', bold=True)
add_bullet('#0B2641 (dark navy background)')
add_bullet('#0D243B (card background)')
add_bullet('#0684F5 (accent blue)')
add_bullet('White / Light gray (text)')
add_para('Typography: Bold sans-serif (Inter/Poppins/Montserrat). Min 24pt for slides.', bold=True)
add_para('Format: 16:9, Max 30MB PDF', bold=True)
add_para('Tone: Confident but not arrogant. Data-driven but narrative. Global product, Africa-Europe beachhead. Show, do not just tell.', bold=True)

doc.save('Eventra_PitchDeck_Content_SoftLanding2026.docx')
print('SUCCESS: Document saved!')
