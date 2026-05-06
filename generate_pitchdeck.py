import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

# Colors
BLUE = RGBColor(0x06, 0x84, 0xF5)
DARK = RGBColor(0x0B, 0x26, 0x41)
GRAY = RGBColor(0x66, 0x66, 0x66)
ORANGE = RGBColor(0xE0, 0x7A, 0x00)
WHITE_ISH = RGBColor(0x33, 0x33, 0x33)

def slide_header(num, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"SLIDE {num}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(2)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(title)
    run2.bold = True
    run2.font.size = Pt(22)
    run2.font.color.rgb = DARK
    p2.paragraph_format.space_after = Pt(8)

def section(label):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE
    run.font.all_caps = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    return p

def visual_note(text):
    p = doc.add_paragraph()
    run = p.add_run(f"\U0001F3A8 VISUAL: {text}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = ORANGE
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

def speaker_note(text):
    p = doc.add_paragraph()
    run = p.add_run(f"\U0001F399 SPEAKER NOTES: ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run2 = p.add_run(text)
    run2.font.size = Pt(9)
    run2.font.color.rgb = GRAY
    run2.italic = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

def divider():
    p = doc.add_paragraph()
    run = p.add_run("\u2500" * 80)
    run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    run.font.size = Pt(6)

# ================================================================
# TITLE PAGE
# ================================================================

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(40)
run = title.add_run("EVENTRA")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Pitch Deck Content Guide")
run.font.size = Pt(16)
run.font.color.rgb = DARK

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run("Soft Landing Provence Africa Connect 2026")
run.font.size = Pt(12)
run.font.color.rgb = GRAY

sub3 = doc.add_paragraph()
sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub3.paragraph_format.space_after = Pt(8)
run = sub3.add_run("10-Slide Guy Kawasaki Format  |  Max 30MB PDF")
run.font.size = Pt(10)
run.font.color.rgb = GRAY

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = note.add_run("This document contains the full text content, speaker notes, and visual direction for each slide.\nUse this to build the final presentation in PowerPoint, Google Slides, or Canva.")
run.font.size = Pt(10)
run.font.color.rgb = GRAY
run.italic = True
note.paragraph_format.space_after = Pt(40)

# ================================================================
# SLIDE 1 — TITLE
# ================================================================
divider()
slide_header(1, "EVENTRA \u2014 The Event Platform Bridging Africa & Europe")

section("On-Screen Content")
body("EVENTRA")
body("The Event Platform Bridging Africa & Europe")
body("")
bullet("Professional Event Management  \u2022  B2B Matchmaking  \u2022  Trilingual (EN/FR/AR)")
body("")
body("Soft Landing Provence Africa Connect 2026")
body("eventra.cloud")

section("Tagline Options (pick one)")
bullet("\"Where Africa meets Europe \u2014 one event at a time.\"")
bullet("\"Build. Connect. Bridge.\"")
bullet("\"Events without borders.\"")

visual_note("Eventra logo centered. Dark background (#0B2641). Blue accent (#0684F5). Subtle world map or Africa-Europe connection graphic. Bottom: logos of SEKETAK, RedStart Tunisia, GITS.")

speaker_note("We are Eventra \u2014 a trilingual event management platform built in Africa, for the Africa-Europe corridor. We help organizers create beautiful event pages, sell tickets, manage attendees, and most importantly \u2014 connect the right people through AI-powered B2B matchmaking. We are applying to Soft Landing to establish our European base in Aix-Marseille-Provence.")

# ================================================================
# SLIDE 2 — PROBLEM
# ================================================================
divider()
slide_header(2, "The Problem")

section("Headline")
body("Africa's event industry is booming \u2014 but the tools weren't built for it.")

section("Key Points")
bullet("$910M+ Middle East & Africa event software market in 2025, growing at 14.8% CAGR \u2014 yet no major platform is built for this region.")
bullet("Western platforms (Eventbrite, Hopin, Bizzabo) are English-only, expensive ($500-$18,000/yr), and ignore African realities: multilingual audiences, B2B trade missions, institutional events.")
bullet("Event organizers in Africa are forced to stitch together 4-5 separate tools: one for registration, one for the website, one for emails, one for networking, one for ticketing.")
bullet("The Africa-Europe business corridor lacks a digital infrastructure for events \u2014 trade missions, investment summits, and innovation forums still rely on spreadsheets and WhatsApp groups.")

section("The Gap")
body("There is no affordable, multilingual, all-in-one event platform designed for the Africa-Europe-MENA corridor.")

visual_note("Split layout. Left side: icons of fragmented tools (Eventbrite + Mailchimp + Wix + Excel + WhatsApp) in a messy cluster. Right side: map highlighting Africa-Europe corridor with a 'gap' icon. Use stats as callout boxes.")

speaker_note("The event tech market in the Middle East and Africa alone is worth 910 million dollars and growing fast. But all the major platforms \u2014 Eventbrite, Hopin, Bizzabo \u2014 were built for Western markets. They are English-first, expensive, and don't understand how events work in Africa. Organizers here run trade missions in French and Arabic, they need B2B matchmaking, and they can't afford 18,000 dollars a year for Bizzabo. They end up using 5 different tools held together with WhatsApp. That's the gap Eventra fills.")

# ================================================================
# SLIDE 3 — SOLUTION
# ================================================================
divider()
slide_header(3, "Our Solution: Eventra")

section("Headline")
body("One platform. Every event. Three languages.")

section("Core Pillars (show as 4 columns or icons)")

body("\U0001F3A8 No-Code Design Studio")
bullet("Drag-and-drop event page builder with 15+ blocks (hero, speakers, agenda, sponsors, gallery, FAQ, countdown...)")
bullet("Fully branded landing pages in minutes \u2014 no designer needed")
body("")

body("\U0001F91D AI-Powered B2B Matchmaking")
bullet("Connects attendees based on profiles, interests, and business goals")
bullet("Structured meeting scheduling within events")
body("")

body("\U0001F30D Trilingual by Architecture")
bullet("English, French, Arabic (with full RTL support) built into the core \u2014 not a translation layer")
bullet("Designed for the Africa-Europe-MENA corridor")
body("")

body("\U0001F4CB All-in-One Management")
bullet("Registration forms, ticketing, speaker management, session scheduling, attendee tracking, enterprise email campaigns \u2014 unified")

visual_note("4-column layout with icons for each pillar. Clean, modern design. Optionally show a mini-screenshot of the Design Studio below. Use Eventra brand colors.")

speaker_note("Eventra solves this with four core pillars. First, a no-code Design Studio \u2014 think of it as Wix for events. Organizers drag and drop 15 different blocks to build beautiful event pages in minutes. Second, AI-powered B2B matchmaking that turns every event into a business connection engine. Third, it's trilingual by architecture, not as an afterthought \u2014 English, French, and Arabic including full right-to-left support. And fourth, everything is unified: registration, ticketing, speakers, sessions, emails, attendee management \u2014 one platform.")

# ================================================================
# SLIDE 4 — PRODUCT DEMO
# ================================================================
divider()
slide_header(4, "The Product")

section("Headline")
body("See it in action \u2014 eventra.cloud")

section("Screenshots to Include (4 key screens)")
bullet("1. Design Studio Editor \u2014 showing the drag-and-drop block interface with the live preview")
bullet("2. Published Event Landing Page \u2014 a real event page (e.g., GITS or IPDAYS) showing hero, speakers, agenda")
bullet("3. B2B Matchmaking / Attendee Dashboard \u2014 showing the networking and meeting scheduling interface")
bullet("4. Organizer Dashboard \u2014 showing event management: attendees list, analytics, forms, email campaigns")

section("Key Callouts (overlay on screenshots)")
bullet("\"15+ customizable blocks\"")
bullet("\"Drag & drop \u2014 zero code\"")
bullet("\"AI-powered attendee matching\"")
bullet("\"Enterprise email campaigns\"")
bullet("\"Real-time analytics\"")

visual_note("This slide is primarily VISUAL. Use 4 screenshots arranged in a 2x2 grid or a staggered/floating mockup layout. Take fresh screenshots from eventra.cloud. Use device mockups (laptop + phone) to make it look polished. Minimal text \u2014 let the product speak.")

speaker_note("Here's Eventra in action. On the top left, our Design Studio where organizers build their event page by dragging blocks. Top right, a published landing page from one of our real events. Bottom left, the B2B matchmaking dashboard. Bottom right, the organizer's management view \u2014 attendees, analytics, email campaigns. All of this is live at eventra.cloud and has been used across 4 international events.")

# ================================================================
# SLIDE 5 — BUSINESS MODEL
# ================================================================
divider()
slide_header(5, "Business Model")

section("Headline")
body("Tiered model: grow with your organizers")

section("Revenue Streams (show as a tiered diagram)")

body("\U0001F7E2 FREE TIER")
bullet("Basic event creation for small events")
bullet("Goal: adoption and platform discovery")
body("")

body("\U0001F535 PAID PLANS")
bullet("Monthly or per-event subscriptions")
bullet("Unlocks: full Design Studio, custom branding, analytics, B2B matchmaking, enterprise emails")
bullet("Target: event organizers, agencies, trade bodies")
body("")

body("\U0001F7E1 ENTERPRISE PACKAGES")
bullet("Custom pricing for large institutional events")
bullet("Dedicated setup, white-label options, on-site support")
bullet("Target: chambers of commerce, ministries, summit organizers")
body("")

body("\U0001F534 TRANSACTION FEES")
bullet("Commission on paid ticket sales processed through Eventra")
body("")

section("Key Message")
body("Pricing is being validated with early adopters during our pilot-to-commercial transition in 2026. Our go-to-market starts with institutional organizers who run recurring annual events \u2014 high retention, predictable revenue.")

visual_note("Show as a pyramid or tiered diagram \u2014 Free at the base (widest), Paid Plans in the middle, Enterprise at the top. Transaction fees as a horizontal bar across all tiers. Use green/blue/gold colors for the tiers. Keep it clean and visual.")

speaker_note("Our model is tiered. A free tier for small events drives adoption. Paid plans unlock the full platform for professional organizers. Enterprise packages serve large institutional events with custom setup and white-label options. And we take a commission on every paid ticket sold through the platform. We are validating pricing now with early adopters. Our first commercial clients will be institutional organizers who run the same events every year \u2014 that means recurring, predictable revenue.")

# ================================================================
# SLIDE 6 — TRACTION
# ================================================================
divider()
slide_header(6, "Traction & Validation")

section("Headline")
body("Battle-tested across 3 countries, 4 events, 1,200+ users")

section("Key Metrics (show as big numbers)")
body("4 \u2014 International events hosted")
body("3 \u2014 Countries (Egypt, C\u00f4te d'Ivoire, Tunisia)")
body("1,200+ \u2014 Registered user accounts")
body("250+ \u2014 International leaders at GITS Cairo")

section("Events Timeline")
bullet("MAY 2025 \u2014 GITS Cairo (Marriott Zamalek) \u2014 170+ global leaders \u2014 With Egypt's Ministry of International Cooperation & GAFI")
bullet("JUNE 2025 \u2014 GITS Abidjan (Azala\u00ef Hotel) \u2014 Francophone Africa expansion")
bullet("NOV 2025 \u2014 IPDAYS x GITS Tunis (Radisson Hotel) \u2014 With RedStart Tunisia \u2014 Major Tunisian press coverage")
bullet("2026 \u2014 Commercial launch & European market entry via Soft Landing")

section("Institutional Validation")
bullet("Egypt's Ministry of International Cooperation")
bullet("GAFI (General Authority for Investment and Free Zones)")
bullet("RedStart Tunisia (800+ startups supported)")
bullet("Press: La Presse de Tunisie, Managers Magazine, Web Manager Center, Tekiano")

visual_note("Big bold numbers on the left (4 events, 3 countries, 1,200+ users, 250+ leaders). Timeline on the right with event logos/photos. Bottom strip: partner logos (Ministry, GAFI, RedStart, media logos). Use real event photos if available.")

speaker_note("We are not a concept \u2014 we are live and tested. Eventra has powered 4 international events across 3 African countries. GITS Cairo brought 170 global trade leaders to the Marriott Zamalek, organized with Egypt's Ministry of International Cooperation. We expanded to Abidjan for Francophone Africa. And IPDAYS in Tunis with RedStart Tunisia got major press coverage. We have 1,200 registered users and institutional validation from government bodies. Now we are ready to bring this to Europe.")

# ================================================================
# SLIDE 7 — MARKET OPPORTUNITY
# ================================================================
divider()
slide_header(7, "Market Opportunity")

section("Headline")
body("$9.9B global market. Africa is the fastest-growing segment.")

section("Market Data")
bullet("Global event management software market: $9.9B in 2025 \u2192 $32.6B by 2034 (14.2% CAGR)")
bullet("Middle East & Africa segment: $910M in 2025 \u2192 $3.3B by 2034 (14.8% CAGR) \u2014 the fastest-growing region")
bullet("Europe: Major secondary market, Marseille is the #1 Euro-African business gateway")

section("Our Addressable Market")
bullet("TAM: $910M \u2014 Full MEA event software market")
bullet("SAM: $150-200M \u2014 B2B/institutional events in Africa + Euro-African corridor")
bullet("SOM: $2-5M \u2014 Institutional event organizers in our direct network (GITS, RedStart, trade missions) within 3-5 years")

section("Why Marseille")
bullet("France's #1 gateway to Africa and the Mediterranean")
bullet("Home to Emerging Valley, the leading Euro-African tech summit")
bullet("Soft Landing program gives direct access to this ecosystem")

visual_note("Concentric circles (TAM > SAM > SOM) on the left. Market growth chart on the right showing the $910M to $3.3B trajectory. Map highlighting Marseille as the bridge point. Source citations at the bottom.")

speaker_note("The global event software market is 9.9 billion dollars and growing to 32 billion by 2034. The Middle East and Africa segment is the fastest growing at nearly 15% per year, reaching 3.3 billion. But there is no dominant platform built for this region. Our addressable market is the 150 to 200 million dollar B2B and institutional events segment in the Africa-Europe corridor. And Marseille is the perfect launchpad \u2014 it's France's number one gateway to Africa, home to Emerging Valley, and exactly where Soft Landing gives us access.")

# ================================================================
# SLIDE 8 — TEAM
# ================================================================
divider()
slide_header(8, "The Team")

section("Headline")
body("Built by a team that lives at the crossroads of Africa and Europe")

section("Assem Kamel \u2014 CEO & Co-Founder")
bullet("B2B strategy, partnerships, business development")
bullet("Co-founder of GITS (Global Investment and Trade Summit)")
bullet("Co-founder of The LEE Experience (Lebanon)")
bullet("Certified Business Coach (International Training Centre of the ILO)")
bullet("15+ years in international trade across MENA and Africa")

section("Douja Gharbi \u2014 Co-Founder & Strategic Partnerships")
bullet("CEO of RedStart Tunisia \u2014 one of North Africa's leading startup accelerators")
bullet("30+ years in entrepreneurship and economic development")
bullet("800+ projects supported, hundreds of entrepreneurs mentored")
bullet("ESIL Angel for Tunisia (European program)")
bullet("Partnerships with UNIDO, INVESTMED, Startup Tunisia")

section("Oussama Lamine \u2014 CTO & Co-Founder")
bullet("Full-stack developer and product architect")
bullet("Built Eventra end-to-end: React/TypeScript, Supabase, AI integrations")
bullet("Design Studio engine, B2B matchmaking, i18n architecture")
bullet("Sole technical lead \u2014 building engineering team is a growth priority")

section("Why This Team Wins")
bullet("Assem brings the clients (GITS network across 3 countries)")
bullet("Douja brings the ecosystem (800+ startups, European institutional partners)")
bullet("Oussama brings the product (built and shipped a full platform)")
bullet("Together: Business + Ecosystem + Technology = the Africa-Europe bridge")

visual_note("Professional headshots of each founder in a row. Name + title below each. Key credentials as bullet points or badge-style callouts. Bottom: 'Why This Team Wins' as a connecting bar showing the trifecta. Include logos: SEKETAK, RedStart, GITS, LEE Experience.")

speaker_note("Our team is our superpower. Assem is a B2B veteran who built GITS into an international trade summit across 3 countries \u2014 he brings the clients. Douja has 30 years of experience and runs RedStart Tunisia, one of North Africa's top accelerators \u2014 she brings the ecosystem and European institutional connections. And Oussama built the entire Eventra platform solo, from the Design Studio to the AI matchmaking \u2014 he brings the technology. Together, we are business, ecosystem, and technology \u2014 the exact combination needed to bridge Africa and Europe.")

# ================================================================
# SLIDE 9 — COMPETITION
# ================================================================
divider()
slide_header(9, "Competitive Landscape")

section("Headline")
body("They built for the West. We built for the bridge.")

section("Comparison Table")
body("Feature           | Eventbrite  | Hopin       | Bizzabo     | EVENTRA")
body("\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500")
body("Multilingual       | \u274C English   | \u274C English   | \u274C English   | \u2705 EN/FR/AR + RTL")
body("B2B Matchmaking    | \u274C No        | \u26A0 Basic      | \u26A0 Add-on     | \u2705 AI-powered, native")
body("No-Code Builder    | \u274C Templates | \u274C Templates | \u26A0 Limited    | \u2705 15+ drag-drop blocks")
body("Africa Focus       | \u274C No        | \u274C No        | \u274C No        | \u2705 Built for Africa-EU")
body("Pricing            | % per ticket | $83-667/mo  | $18,000/yr+ | Free + affordable plans")
body("All-in-One         | \u274C Ticketing | \u26A0 Virtual    | \u2705 Yes       | \u2705 Yes")

section("Our Edge")
bullet("Only platform with native trilingual support (EN/FR/AR with RTL)")
bullet("Only platform with AI B2B matchmaking built-in from day one")
bullet("Purpose-built for Africa-Europe institutional events and trade missions")
bullet("10x more affordable than Bizzabo, more powerful than Eventbrite")

visual_note("Clean comparison table or matrix chart. Use checkmarks (green), X marks (red), and warning signs (yellow). Eventra column highlighted in blue. Keep it scannable. Competitor logos in the column headers.")

speaker_note("Let me be clear about what sets us apart. Eventbrite is a ticketing tool \u2014 no page builder, no matchmaking, English only. Hopin is for virtual events and costs up to 667 dollars a month. Bizzabo is enterprise-only at 18,000 dollars a year minimum. None of them speak French or Arabic. None of them are built for African realities. Eventra is the only platform that combines a no-code design studio, AI matchmaking, trilingual support, and all-in-one event management \u2014 at an accessible price point. We didn't just add features. We built for a market they ignore.")

# ================================================================
# SLIDE 10 — ASK & ROADMAP
# ================================================================
divider()
slide_header(10, "The Ask & Roadmap")

section("Headline")
body("From Soft Landing to European Market Leader")

section("What We Need from the Program")
bullet("\u2705 Legal entity setup in Aix-Marseille-Provence for European operations")
bullet("\u2705 B2B introductions to event organizers, chambers of commerce, and institutional partners in France")
bullet("\u2705 Investor access for our pre-seed/seed round")
bullet("\u2705 Mentorship on European market entry strategy, GDPR compliance, and pricing")
bullet("\u2705 Participation in Le Grand Bain and Emerging Valley to showcase Eventra")

section("Roadmap")
body("H2 2026 \u2014 ESTABLISH")
bullet("European entity setup via Soft Landing")
bullet("First paid European clients (institutional events)")
bullet("Pricing validation and commercial launch")
body("")
body("2027 \u2014 GROW")
bullet("10-15 paying clients across Africa and Europe")
bullet("\u20AC15K-25K first-year revenue")
bullet("Hire first European team member (sales/partnerships)")
bullet("Pre-seed fundraising round")
body("")
body("2028 \u2014 SCALE")
bullet("50+ paying clients")
bullet("Seed round fundraising")
bullet("Expand to additional languages (Portuguese, Spanish)")
bullet("Become the default event platform for the Africa-Europe corridor")

section("Call to Action")
body("We are not asking you to bet on an idea.")
body("We are asking you to back a product that already works,")
body("a team that already delivers,")
body("and a market that is already growing.")
body("")
body("eventra.cloud  |  team@seketak-eg.com")

visual_note("Left side: 3-step vertical roadmap (Establish > Grow > Scale) with year labels and milestones. Right side: 'What we need' as a checklist. Bottom: bold call-to-action text. Eventra logo and contact info. End on confidence, not desperation.")

speaker_note("Here is what we need from Soft Landing. Help us set up our European legal entity. Introduce us to institutional event organizers in France. Connect us with investors for our first funding round. And give us a stage at Le Grand Bain and Emerging Valley. Our roadmap is clear: establish in 2026, grow in 2027, scale in 2028. We are not asking you to bet on a pitch. Eventra is live. It has powered 4 international events across 3 countries. We have 1,200 users and institutional backing from government bodies. We are asking you to help us bring this bridge to Europe. Thank you.")

# ================================================================
# SAVE
# ================================================================
output_path = os.path.join(
    os.environ.get('USERPROFILE', ''),
    'Desktop', 'Agent X', 'EVENTRA X', 'Eventra Code V1.1', 'eventra',
    'Eventra_PitchDeck_Content_SoftLanding2026.docx'
)
doc.save(output_path)
print(f'SUCCESS: {output_path}')
